"""
NAIdo Doc Reader — Lecture et extraction de contenu
Supporte: Word (.docx), PDF, Excel (.xlsx)
"""
from pathlib import Path
import re

def extraire_texte_docx(file_path: str) -> str:
    """Extrait le texte d'un fichier Word"""
    try:
        from docx import Document
        doc = Document(file_path)
        texte = []
        for para in doc.paragraphs:
            if para.text.strip():
                texte.append(para.text.strip())
        # Tables aussi
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texte.append(cell.text.strip())
        return "\n".join(texte)
    except Exception as e:
        return f"Erreur lecture Word: {e}"

def extraire_texte_pdf(file_path: str) -> str:
    """Extrait le texte d'un PDF"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        texte = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                texte.append(t)
        return "\n".join(texte)
    except Exception as e:
        return f"Erreur lecture PDF: {e}"

def extraire_texte_xlsx(file_path: str) -> str:
    """Extrait le texte d'un Excel"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        texte = []
        for sheet in wb.worksheets:
            texte.append(f"=== {sheet.title} ===")
            for row in sheet.iter_rows(max_row=100):
                vals = [str(cell.value) for cell in row if cell.value is not None]
                if vals:
                    texte.append(" | ".join(vals))
        return "\n".join(texte)
    except Exception as e:
        return f"Erreur lecture Excel: {e}"

def extraire_contenu(file_path: str) -> dict:
    """Extrait le contenu selon le type de fichier"""
    path = Path(file_path)
    ext = path.suffix.lower()
    
    if ext in ('.docx', '.doc'):
        texte = extraire_texte_docx(file_path)
    elif ext == '.pdf':
        texte = extraire_texte_pdf(file_path)
    elif ext in ('.xlsx', '.xls'):
        texte = extraire_texte_xlsx(file_path)
    else:
        return {"texte": "", "mots_cles": [], "nb_mots": 0}
    
    # Nettoyer
    texte = re.sub(r'\s+', ' ', texte).strip()
    texte_tronque = texte[:8000]  # Max 8000 chars pour l'IA
    
    # Extraire mots-clés simples (les mots longs fréquents)
    mots = re.findall(r'\b[a-zA-ZÀ-ÿ]{5,}\b', texte.lower())
    freq = {}
    for m in mots:
        freq[m] = freq.get(m, 0) + 1
    # Top 20 mots-clés (exclus les mots vides)
    STOP = {'cette','dans','pour','avec','sont','être','avoir','faire',
            'plus','mais','tout','bien','dont','aussi','comme','leur'}
    mots_cles = [m for m,_ in sorted(freq.items(),key=lambda x:-x[1]) 
                 if m not in STOP][:20]
    
    return {
        "texte": texte_tronque,
        "texte_complet": texte,
        "mots_cles": mots_cles,
        "nb_mots": len(texte.split()),
        "nb_pages": texte.count('\n\n') + 1,
    }

def info_fichier(file_path: str) -> dict:
    """Infos basiques sur le fichier"""
    path = Path(file_path)
    return {
        "taille_kb": round(path.stat().st_size / 1024, 1) if path.exists() else 0,
        "extension": path.suffix.lower(),
        "nom": path.name,
    }
